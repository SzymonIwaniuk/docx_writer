import datetime
import os

from docx import Document

from src.domain.services import change_item_contract


def test_change_item_contract_content() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "take_id": "K123",
        "take_item": "Laptop Asus 12345AB",
        "take_qty": "1",
        "give_id": "K321",
        "give_item": "Laptop Dell 12345AB",
        "give_qty": "1",
        "date": "2025-08-11",
    }

    creation_path = change_item_contract(**data)

    try:
        doc = Document(creation_path)

        content = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += "\n" + cell.text

        for value in data.values():
            assert value in content

    finally:
        # Delete file
        os.remove(creation_path)


def test_change_item_contract_save_path() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "take_id": "K123",
        "take_item": "Laptop Asus 12345AB",
        "take_qty": "1",
        "give_id": "K321",
        "give_item": "Laptop Dell 12345AB",
        "give_qty": "1",
        "date": "2025-08-11",
    }

    creation_path = change_item_contract(**data)

    try:
        assert os.path.exists(creation_path)
        assert creation_path.endswith(".pdf")
        assert data["borrower"].replace(" ", "_") in creation_path

    finally:
        # Delete file
        os.remove(creation_path)


def test_change_item_contract_fill_with_today_date() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "take_id": "K123",
        "take_item": "Laptop Asus 12345AB",
        "take_qty": "1",
        "give_id": "K321",
        "give_item": "Laptop Dell 12345AB",
        "give_qty": "1",
        "date": None,
    }

    creation_path = change_item_contract(**data)

    try:
        doc = Document(creation_path)
        content = "\n".join([p.text for p in doc.paragraphs])

        today = datetime.datetime.today().strftime("%d-%m-%Y")
        assert today in content

    finally:
        # Delete file
        # os.remove(creation_path)
        pass
