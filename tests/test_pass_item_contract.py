import datetime
import os

from docx import Document

from src.domain.services import pass_item_contract


def test_pass_item_contract_content() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "id": "K123",
        "item": "Laptop Dell 12345AB",
        "quantity": "1",
        "date": "2025-08-11",
    }

    creation_path = pass_item_contract(**data)

    try:
        doc = Document(creation_path)

        content = "\n".join([p.text for p in doc.paragraphs])

        # Get data from ceils
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += "\n" + cell.text

        for value in data.values():
            assert value in content

    finally:
        # Delete file
        os.remove(creation_path)


def test_pass_item_contract_save_path() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "id": "K123",
        "item": "Laptop Dell 12345AB",
        "quantity": "1",
        "date": "2025-08-11",
    }

    creation_path = pass_item_contract(**data)
    try:
        assert os.path.exists(creation_path)
        assert creation_path.endswith(".pdfd")
        assert data["borrower"].replace(" ", "_") in creation_path

    finally:
        # Delete file
        os.remove(creation_path)


def test_pass_item_contract_fill_with_today_date() -> None:
    data = {
        "it_worker": "Szymon Iwaniuk",
        "borrower": "Mike Wazowski",
        "id": "K123",
        "item": "Laptop Dell 12345AB",
        "quantity": "1",
    }

    creation_path = pass_item_contract(**data)

    try:
        doc = Document(creation_path)
        content = "\n".join([p.text for p in doc.paragraphs])

        # Today date in YYYY-MM-DD format
        date = datetime.datetime.today().strftime("%d-%m-%Y")

        assert date in content

    finally:
        # Delete file
        os.remove(creation_path)
