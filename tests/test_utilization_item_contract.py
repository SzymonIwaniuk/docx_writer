import datetime
import os

import pytest
from docx import Document

from src.domain.services import utilization_items_contract


def test_utilization_items_contract_content() -> None:
    data = [
        {"id": "K123", "name": "laptop1", "inventarization_num": "1077", "date": "2025-08-01"},
        {"id": "K133", "name": "laptop2", "inventarization_num": "1078", "date": "2025-08-02"},
        {"id": "K211", "name": "laptop3", "inventarization_num": "1079", "date": "2025-08-03"},
    ]

    participants = ["Szymon Iwaniuk", "Mike Wazowski"]

    creation_path = utilization_items_contract(items=data, participants=participants)

    try:
        doc = Document(creation_path)

        content = "\n".join([p.text for p in doc.paragraphs])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += "\n" + cell.text

        for item in data:
            for value in item.values():
                assert value in content

    finally:
        os.remove(creation_path)


def test_utilization_items_contract_save_path() -> None:
    data = [
        {"id": "K123", "name": "laptop1", "inventarization_num": "1077", "date": "2025-08-01"},
        {"id": "K133", "name": "laptop2", "inventarization_num": "1078", "date": "2025-08-02"},
        {"id": "K211", "name": "laptop3", "inventarization_num": "1079", "date": "2025-08-03"},
    ]

    participants = ["Szymon Iwaniuk", "Mike Wazowski"]

    try:
        creation_path = utilization_items_contract(items=data, participants=participants)

        assert os.path.exists(creation_path)
        assert creation_path.endswith(".pdf")
        assert "Utylizacja_sprzetu_" in os.path.basename(creation_path)

    finally:
        os.remove(creation_path)


def test_utilization_items_contract_fill_with_today_date() -> None:
    data = [
        {"id": "K123", "name": "laptop1", "inventarization_num": "1077", "date": "2025-08-01"},
        {"id": "K133", "name": "laptop2", "inventarization_num": "1078", "date": "2025-08-02"},
        {"id": "K211", "name": "laptop3", "inventarization_num": "1079", "date": "2025-08-03"},
    ]

    participants = ["Szymon Iwaniuk", "Mike Wazowski"]

    try:
        creation_path = utilization_items_contract(items=data, participants=participants)

        doc = Document(creation_path)

        content = "\n".join([p.text for p in doc.paragraphs])
        today = datetime.datetime.today().strftime("%d-%m-%Y")

        assert today in content

    finally:
        os.remove(creation_path)
