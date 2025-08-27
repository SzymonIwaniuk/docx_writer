import datetime
import os
from typing import List

from docx import Document

from src.domain.helper import resource_path, save_as_pdf, set_cell_border


def pass_item_contract(it_worker: str, borrower: str, id: str, item: str, quantity: str, date=None) -> str:
    """
    Generates a Word document contract for passing IT equipment to a borrower and convert to pdf.

    This function loads a predefined Word template, replaces placeholders with provided values,
    and saves the filled document converted to pdf to the C:/docx_wrtier/attachments with a filename
    based on the borrower's name and date.

    Args:
        it_worker (str): Name of the IT worker handing over the item.
        borrower (str): Name of the person receiving the item.
        date (Optional[str]): Date of the handover. If None, defaults to today's date in DD-MM- format.
        id (str): Identifier for the transaction or item.
        item (str): Description of the item being handed over.
        quantity (str): Quantity of the item being handed over.

    Returns:
    str: Path of saved pdf file.
    """

    template_path = resource_path(os.path.join("src", "templates", "pass_item_template.docx"))
    doc = Document(template_path)

    if date is None:
        date = datetime.datetime.today().strftime("%d-%m-%Y")

    replacements = {
        "{{it_worker}}": it_worker,
        "{{borrower}}": borrower,
        "{{date}}": date,
        "{{id}}": id,
        "{{item}}": item,
        "{{quantity}}": quantity,
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    filename = f"Przekazanie_sprzetu_{borrower.replace(' ', '_')}_{date}.docx"
    save_path = os.path.join(r"C:\docx_writer\attachments", filename)

    # Save document
    doc.save(save_path)

    # Conversion to pdf
    pdf_path = save_as_pdf(save_path)
    os.remove(save_path)
    return pdf_path


def change_item_contract(
    it_worker: str,
    borrower: str,
    take_id: str,
    take_item: str,
    take_qty: str,
    give_id: str,
    give_item: str,
    give_qty: str,
    date=None,
) -> str:
    """
    Generates a Word document contract for exchanging IT equipment between a borrower and IT worker,
    then converts it to PDF.

    This function loads a predefined Word template, replaces placeholders with provided values,
    and saves the filled document as a PDF in `C:/docx_writer/attachments`.
    The filename is based on the borrower's name and the date.

    Args:
        it_worker (str): Name of the IT worker facilitating the exchange.
        borrower (str): Name of the person exchanging items.
        take_id (str): Identifier for the item being taken.
        take_item (str): Description of the item being taken.
        take_qty (str): Quantity of the item being taken.
        give_id (str): Identifier for the item being given.
        give_item (str): Description of the item being given.
        give_qty (str): Quantity of the item being given.
        date (Optional[str]): Date of the exchange. Defaults to today's date (DD-MM-YYYY).

    Returns:
        str: Path of the saved PDF file.
    """

    template_path = resource_path(os.path.join("src", "templates", "change_item_template.docx"))
    doc = Document(template_path)

    if date is None:
        date = datetime.datetime.today().strftime("%d-%m-%Y")

    replacements = {
        "{{it_worker}}": it_worker,
        "{{borrower}}": borrower,
        "{{take_id}}": take_id,
        "{{take_item}}": take_item,
        "{{take_qty}}": take_qty,
        "{{give_id}}": give_id,
        "{{give_item}}": give_item,
        "{{give_qty}}": give_qty,
        "{{date}}": date,
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    filename = f"Wymiana_sprzetu_{borrower.replace(' ', '_')}_{date}.docx"
    save_path = os.path.join(r"C:\docx_writer\attachments", filename)

    # Save document
    doc.save(save_path)

    # Conversion to pdf
    pdf_path = save_as_pdf(save_path)
    os.remove(save_path)
    return pdf_path


def utilization_items_contract(
    items: list,
    participants: List[str],
    date=None,
) -> str:
    """
    Generates a Word document contract for the utilization of IT equipment and converts it to PDF.

    This function loads a predefined Word template, fills it with participant information and
    a table of items to be utilized, then saves the completed document as a PDF in
    `C:/docx_writer/attachments`. The filename is based on the date.

    Args:
        items (list[dict]): List of items for utilization. Each item dictionary should contain:
            - id (str): Identifier of the item.
            - name (str): Name or description of the item.
            - inventarization_num (str): Inventory number of the item.
            - date (str): Associated date for the item.
        participants (List[str]): Names of participants overseeing the utilization process.
        date (Optional[str]): Date of the utilization. Defaults to today's date (DD-MM-YYYY).

    Returns:
        str: Path of the saved PDF file.
    """

    template_path = resource_path(os.path.join("src", "templates", "utilization_items_template.docx"))
    doc = Document(template_path)

    if date is None:
        date = datetime.datetime.today().strftime("%d-%m-%Y")

    participants_section = ""
    for name in participants:
        participants_section += f"{name} " + "\n" + "." * 40 + "\n"

    replacements = {
        "{{participants_section}}": participants_section,
        "{{date}}": date,
    }

    # Replace with provided data
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

    if doc.tables:
        table = doc.tables[0]
        for item in items:
            row = table.add_row()
            row.cells[0].text = item.get("id", "")
            row.cells[1].text = item.get("name", "")
            row.cells[2].text = item.get("inventarization_num", "")
            row.cells[3].text = item.get("date", "")
            for cell in row.cells:
                set_cell_border(cell, size="4", color="000000")

    filename = f"Utylizacja_sprzetu_{date}.docx"
    save_path = os.path.join(r"C:\docx_writer\attachments", filename)

    # Save document
    doc.save(save_path)

    # Conversion to pdf
    pdf_path = save_as_pdf(save_path)
    os.remove(save_path)
    return pdf_path
