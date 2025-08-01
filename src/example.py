import os
from datetime import datetime
from typing import Optional


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell


def create_laptop_borrow_contract(
    it_worker: str,
    borrower: str,
    date: str,
    lp: int,
    item_name: str,
    model_sn: str,
    quantity: int,
    return_date: str,
    template_path: Optional[str] = None,
    save_path: Optional[str] = None,
) -> str:

    # Load template or create new doc
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # Add logo on the top and center it
    logo = doc.add_heading("", 0)
    logo_color = logo.add_run("SaMASZ Sp z.o.o")
    logo_color.font.color.rgb = RGBColor(0, 128, 0)
    logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    top_text = (
        "SaMASZ sp. z o.o. 16-060, Zabłudów, ul Trawiasta 1\n"
        "tel.: +48 85 664 70 31, fax: +48 85 664 7041, e-mail: samasz@samasz.pl\n"
        "KRS:0000011382, Sąd Rejonowy w Białymstoku, XII Wydział Gospodarczy Krajowego Rejestru Sądowego\n"
        "Kapitał zakładowy: 10 000 000zł."
    )

    top_paragraph = doc.add_paragraph()
    top_run = top_paragraph.add_run(top_text)
    font = top_run.font
    font.name = "Arial"
    font.size = Pt(7.5)
    font.bold = True
    font.color.rgb = RGBColor(0, 128, 0)  # Green

    # Spacing
    top_paragraph_format = top_paragraph.paragraph_format
    top_paragraph_format.line_spacing = Pt(14)

    """
    PROTOKÓŁ PRZEKAZANIA 
    SPRZĘTU KOMPUTEROWEGO 
    """

    top_protocol = doc.add_paragraph("")
    top_text_protocol = top_protocol.add_run("PROTOKÓŁ PRZEKAZANIA")
    top_text_protocol.font.color.rgb = RGBColor(0, 0, 0)
    top_text_protocol.bold = True
    top_text_protocol.font.name = "Times New Roman"
    top_text_protocol.font.size = Pt(14.5)
    top_text_protocol.font.bold = True
    top_protocol.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bot_protocol = doc.add_paragraph("")
    bot_text_protocol = bot_protocol.add_run("SPRZĘTU KOMPUTEROWEGO")
    bot_text_protocol.font.color.rgb = RGBColor(0, 0, 0)
    bot_text_protocol.font.name = "Times New Roman"
    bot_text_protocol.font.size = Pt(11)
    bot_protocol.alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
    Zabłudów, dnia 30.07.2025 
    """
    city_and_date = doc.add_paragraph("")
    city_and_date_text = city_and_date.add_run(f"Zabłudów, dnia {date}")
    city_and_date_text.font.color.rgb = RGBColor(0, 0, 0)
    city_and_date_text.font.name = "Times New Roman"
    city_and_date_text.font.size = Pt(11)
    city_and_date.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Empty line
    doc.add_paragraph("")

    """
    Przekazujący –               Piotr Kozłowski DATE przekazuje  
    """

    IT_worker = doc.add_paragraph("")
    IT_worker_text = IT_worker.add_run(f"Przekazujący –               {it_worker} {date} przekazuje")
    IT_worker_text.font.color.rgb = RGBColor(0, 0, 0)
    IT_worker_text.font.name = "Times New Roman"
    IT_worker_text.font.bold = True
    IT_worker_text.font.size = Pt(12)

    # Empty line
    doc.add_paragraph("")

    """
    Odbierający –                Piotr Lewczuk DATE odbiera,  
    """

    Borrower = doc.add_paragraph("")
    Borrower_text = Borrower.add_run(f"Odbierający –               {borrower} {date} odbiera,")
    Borrower_text.font.color.rgb = RGBColor(0, 0, 0)
    Borrower_text.font.name = "Times New Roman"
    Borrower_text.font.bold = True
    Borrower_text.font.size = Pt(12)

    """
    Odbiera niżej wymieniony przedmiot:
    """

    # Empty
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=2, cols=4)
    table.autofit = True
    headers = ["Lp.", "Przedmiot (Nazwa)", "Nazwa, model, S/N", "Ilość"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header

    placeholders = [f"{lp}", f"{item_name}", f"{model_sn}", f"{quantity}"]
    for i, placeholder in enumerate(placeholders):
        cell = table.cell(1, i)
        cell.text = placeholder
    
    def set_cell_border(cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
     
        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
     
                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
     
                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 8, "val": "single", "color": "FFFFFF"},
                bottom={"sz": 8, "val": "single", "color": "FFFFFF"},
                start={"sz": 8, "val": "single", "color": "FFFFFF"},
                end={"sz": 8, "val": "single", "color": "FFFFFF"},
            )

    # Test logo image
    pic = doc.add_picture("assets/logo.png")
    last = doc.paragraphs[-1]

    # Save path
    if not save_path:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        filename = f"Laptop_Contract_{borrower.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        save_path = os.path.join(desktop, filename)

    # Save document
    doc.save(save_path)
    print(f"Document saved to: {save_path}")
    return save_path


if __name__ == "__main__":
    create_laptop_borrow_contract(
        it_worker="Szymon Iwaniuk",
        borrower="Mike Wazowski",
        lp=1,
        item_name="laptop",
        model_sn="dell",
        quantity=2,
        date="2025-07-31",
        return_date="2025-08-31",
        template_path=None,
        save_path=None,
    )
